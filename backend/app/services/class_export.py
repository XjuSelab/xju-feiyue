"""Render a class's 分组 (groups) into the 课程设计选题一览表 .docx.

Mirrors the layout of the reference workbook shipped with the course
(《XX班_工程项目开发综合实践题目汇总.docx》): a title, a course header line,
then one table whose columns are 序号 / 团队 / 学号 / 姓名 / 学年（课程）论文题目 /
Logo. Each group spans one row per member; the 序号 / 团队 / 论文题目 / Logo cells
are vertically merged across the group's members (as in the template), 论文题目 is
filled from the group intro, and the group logo image is embedded from disk.

Format is aligned to the reference sheet:
  * 横版 A4（29.7×21cm），页边距 上下 1.25″、左右 1″
  * 全文 仿宋（含中文 eastAsia 字形）
  * 标题 三号（16pt），正文 五号（10.5pt）
  * 表格全框线（single, 0.5pt），表头居中

Kept out of services.groups so the python-docx dependency stays isolated to the
one export code path. 班委-only — the route enforces it.
"""

from __future__ import annotations

from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL, WD_ROW_HEIGHT_RULE
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt
from sqlalchemy import select
from sqlalchemy.ext.asyncio import AsyncSession

from app.db.models import Group, GroupMember, StudentClass, User
from app.services.groups import UPLOAD_ROOT

# 指导教师 as printed on the reference sheet; other header fields are class-derived
# or left blank for hand-fill (matches the template's blank 专业 cell).
_TEACHER = "郑炅"
_COURSE = "工程项目开发综合实践"
_COLLEGE = "计算机学院"
_HEADERS = ["序号", "团队", "学号", "姓 名", "学年（课程）论文题目", "Logo"]

# 全文字体 & 字号（对齐参考 docx）——仿宋，标题三号/正文五号。
_FONT = "仿宋"
_TITLE_PT = 16.0  # 三号
_BODY_PT = 10.5  # 五号

# 列宽（英寸）——量自参考 docx 的六列。
_COL_WIDTHS_IN = (0.51, 1.56, 1.48, 1.77, 3.45, 1.30)


def _apply_font(run, *, size_pt: float = _BODY_PT, bold: bool = False) -> None:
    """仿宋（ascii + 中文 eastAsia）+ 指定字号。

    python-docx 的 ``run.font.name`` 只写 ascii/hAnsi 字形，中文字要单独在
    ``w:rFonts@w:eastAsia`` 上指定，否则中文回落到默认宋体。
    """
    run.font.name = _FONT
    run.font.bold = bold
    run.font.size = Pt(size_pt)
    rpr = run._element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:eastAsia"), _FONT)


def _set_document_defaults(doc: Document) -> None:
    """横版 A4 + 页边距，并把 Normal 样式默认字体设为仿宋五号（兜底继承）。"""
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width = Cm(29.7)
    sec.page_height = Cm(21.0)
    sec.left_margin = Inches(1.0)
    sec.right_margin = Inches(1.0)
    sec.top_margin = Inches(1.25)
    sec.bottom_margin = Inches(1.25)

    normal = doc.styles["Normal"]
    normal.font.name = _FONT
    normal.font.size = Pt(_BODY_PT)
    rpr = normal.element.get_or_add_rPr()
    rpr.get_or_add_rFonts().set(qn("w:eastAsia"), _FONT)


def _logo_disk_path(group: Group) -> Path | None:
    """Map a group's logo URL (…/uploads/groups/<gid>/<fname>) to its disk file."""
    if not group.logo:
        return None
    fname = group.logo.rsplit("/", 1)[-1]
    if not fname:
        return None
    path = UPLOAD_ROOT / "groups" / group.id / fname
    return path if path.is_file() else None


async def _load_groups(
    db: AsyncSession, class_id: int
) -> list[tuple[Group, list[tuple[str, str]]]]:
    """(group, [(sid, name), …]) for every live group, ordered as on the sheet."""
    groups = list(
        (
            await db.execute(
                select(Group)
                .where(Group.class_id == class_id, Group.deleted == False)  # noqa: E712
                .order_by(Group.created_at)
            )
        )
        .scalars()
        .all()
    )
    out: list[tuple[Group, list[tuple[str, str]]]] = []
    for g in groups:
        rows = (
            await db.execute(
                select(GroupMember.sid, User.name)
                .join(User, User.sid == GroupMember.sid)
                .where(GroupMember.group_id == g.id)
                # leader first, then join order — the leader heads the team block.
                .order_by((GroupMember.role != "leader"), GroupMember.joined_at, GroupMember.sid)
            )
        ).all()
        out.append((g, [(sid, name) for sid, name in rows]))
    return out


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    center: bool = False,
    size_pt: float = _BODY_PT,
) -> None:
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    para = cell.paragraphs[0]
    if center:
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.text = ""
    run = para.add_run(text)
    _apply_font(run, size_pt=size_pt, bold=bold)


def _merge_down(table, col: int, top_row: int, bottom_row: int):
    """Vertically merge a column across [top_row, bottom_row]; return merged cell."""
    top = table.cell(top_row, col)
    if bottom_row > top_row:
        top = top.merge(table.cell(bottom_row, col))
    return top


def _apply_col_widths(table) -> None:
    """Pin the six column widths on every row (Word honours per-cell widths)."""
    widths = [Inches(w) for w in _COL_WIDTHS_IN]
    table.autofit = False
    table.allow_autofit = False
    for i, w in enumerate(widths):
        table.columns[i].width = w
    for row in table.rows:
        for i, w in enumerate(widths):
            row.cells[i].width = w


async def build_class_groups_docx(
    db: AsyncSession, clazz: StudentClass, class_id: int
) -> bytes:
    """Return the 课程设计选题一览表 .docx bytes for a class's groups."""
    groups = await _load_groups(db, class_id)

    doc = Document()
    _set_document_defaults(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _apply_font(title.add_run("新疆大学本科课程设计选题一览表"), size_pt=_TITLE_PT)

    header = doc.add_paragraph()
    _apply_font(
        header.add_run(
            f"学院：{_COLLEGE}    专业：            课程名：{_COURSE}    "
            f"班级：{clazz.short_name or ''}    指导教师：{_TEACHER}"
        ),
        size_pt=_BODY_PT,
    )

    table = doc.add_table(rows=1, cols=len(_HEADERS))
    table.style = "Table Grid"
    hdr = table.rows[0]
    hdr.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
    hdr.height = Pt(28)
    for i, text in enumerate(_HEADERS):
        _set_cell_text(hdr.cells[i], text, center=True)

    for idx, (group, members) in enumerate(groups, start=1):
        # At least one row per group even when it has no members yet.
        span = max(len(members), 1)
        first = len(table.rows)
        for r in range(span):
            row = table.add_row()
            row.height_rule = WD_ROW_HEIGHT_RULE.AT_LEAST
            row.height = Pt(24)
            if r < len(members):
                sid, name = members[r]
                _set_cell_text(row.cells[2], sid, center=True)
                _set_cell_text(row.cells[3], name, center=True)
        last = len(table.rows) - 1

        _set_cell_text(_merge_down(table, 0, first, last), str(idx), center=True)
        _set_cell_text(_merge_down(table, 1, first, last), group.name, center=True)
        _set_cell_text(_merge_down(table, 4, first, last), group.intro or "")

        logo_cell = _merge_down(table, 5, first, last)
        logo_cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        logo_path = _logo_disk_path(group)
        if logo_path is not None:
            para = logo_cell.paragraphs[0]
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                para.add_run().add_picture(str(logo_path), width=Inches(0.8))
            except Exception:
                # Unreadable / non-image file — leave the cell blank rather than 500.
                pass

    _apply_col_widths(table)

    buf = BytesIO()
    doc.save(buf)
    return buf.getvalue()
