"""/classes/me/missions/* + /classes/me/groups/export tests.

Covers: mission CRUD, the single-active invariant (设为进行中 unsets the prior
one), committee-only writes with cross-class isolation, and the .docx export
(committee 200 + content-type / filename, non-committee 403, member reads).
"""

from __future__ import annotations

from httpx import AsyncClient

DOCX_MIME = "application/vnd.openxmlformats-officedocument.wordprocessingml.document"

COMMITTEE = "20240001001"
MEMBER1 = "20240001002"
MEMBER2 = "20240001003"
ADMIN_MEMBER = "20240001004"
OTHER_COMMITTEE = "20240002001"


async def _create_mission(
    client: AsyncClient, headers: dict[str, str], title: str, **body
) -> dict:
    r = await client.post(
        "/classes/me/missions", json={"title": title, **body}, headers=headers
    )
    assert r.status_code == 201, r.text
    return r.json()


async def _create_group(client: AsyncClient, headers: dict[str, str], name: str) -> str:
    r = await client.post(
        "/classes/me/groups", json={"name": name, "intro": f"{name}的课设选题"}, headers=headers
    )
    assert r.status_code == 201, r.text
    return r.json()["id"]


# --- missions CRUD + permissions -------------------------------------------


async def test_create_mission_committee_only(client: AsyncClient, class_setup, login) -> None:
    m = await _create_mission(client, await committee(login), "第一次分组", description="按兴趣分组")
    assert m["title"] == "第一次分组"
    assert m["isActive"] is True  # active by default


async def test_create_mission_member_forbidden(client: AsyncClient, class_setup, login) -> None:
    r = await client.post(
        "/classes/me/missions", json={"title": "偷偷建"}, headers=await login(MEMBER1)
    )
    assert r.status_code == 403


async def test_mission_list_visible_to_members(client: AsyncClient, class_setup, login) -> None:
    await _create_mission(client, await committee(login), "任务A")
    r = await client.get("/classes/me/missions", headers=await login(MEMBER1))
    assert r.status_code == 200
    assert [m["title"] for m in r.json()] == ["任务A"]


async def test_set_active_is_mutually_exclusive(client: AsyncClient, class_setup, login) -> None:
    headers = await committee(login)
    a = await _create_mission(client, headers, "任务A")  # active
    b = await _create_mission(client, headers, "任务B", active=False)
    assert b["isActive"] is False

    # Activating B must deactivate A.
    r = await client.patch(
        f"/classes/me/missions/{b['id']}", json={"active": True}, headers=headers
    )
    assert r.status_code == 200
    assert r.json()["isActive"] is True

    ordered = (await client.get("/classes/me/missions", headers=headers)).json()
    by_id = {m["id"]: m for m in ordered}
    assert by_id[b["id"]]["isActive"] is True
    assert by_id[a["id"]]["isActive"] is False
    # active mission sorts first
    assert ordered[0]["id"] == b["id"]


async def test_create_active_unsets_prior(client: AsyncClient, class_setup, login) -> None:
    headers = await committee(login)
    a = await _create_mission(client, headers, "任务A")
    b = await _create_mission(client, headers, "任务B")  # active default → unsets A
    ms = {m["id"]: m["isActive"] for m in (await client.get("/classes/me/missions", headers=headers)).json()}
    assert ms[b["id"]] is True
    assert ms[a["id"]] is False


async def test_update_and_delete_mission(client: AsyncClient, class_setup, login) -> None:
    headers = await committee(login)
    m = await _create_mission(client, headers, "旧标题")
    r = await client.patch(
        f"/classes/me/missions/{m['id']}", json={"title": "新标题"}, headers=headers
    )
    assert r.status_code == 200 and r.json()["title"] == "新标题"

    r = await client.delete(f"/classes/me/missions/{m['id']}", headers=headers)
    assert r.status_code == 204
    assert (await client.get("/classes/me/missions", headers=headers)).json() == []


async def test_mission_cross_class_isolation(client: AsyncClient, class_setup, login) -> None:
    m = await _create_mission(client, await committee(login), "本班任务")
    # Other class's committee can't see or touch it (404 — cross-class = unseen).
    other = await login(OTHER_COMMITTEE)
    assert (await client.get("/classes/me/missions", headers=other)).json() == []
    r = await client.patch(f"/classes/me/missions/{m['id']}", json={"title": "x"}, headers=other)
    assert r.status_code == 404


# --- export -----------------------------------------------------------------


async def test_export_committee_downloads_docx(client: AsyncClient, class_setup, login) -> None:
    # A couple of groups so the table has real rows.
    await _create_group(client, await login(MEMBER1), "飞跃组")
    await _create_group(client, await login(MEMBER2), "启航组")

    r = await client.get("/classes/me/groups/export", headers=await committee(login))
    assert r.status_code == 200
    assert r.headers["content-type"].startswith(DOCX_MIME)
    cd = r.headers["content-disposition"]
    assert "attachment" in cd
    # RFC 5987 filename* carries the (percent-encoded) 班级简称 basename.
    assert "filename*=UTF-8''" in cd
    # A .docx is a zip — check the PK magic bytes and non-empty payload.
    assert r.content[:2] == b"PK"

    # Parse it back and assert the reference table shape.
    from io import BytesIO

    from docx import Document

    doc = Document(BytesIO(r.content))
    assert len(doc.tables) == 1
    table = doc.tables[0]
    header = [c.text.strip() for c in table.rows[0].cells]
    assert header == ["序号", "团队", "学号", "姓 名", "学年（课程）论文题目", "Logo"]

    all_text = "\n".join(c.text for row in table.rows for c in row.cells)
    # Team names, the leader's sid/name, and the intro (论文题目 列) all present.
    assert "飞跃组" in all_text and "启航组" in all_text
    assert MEMBER1 in all_text and "成员乙" in all_text
    assert "飞跃组的课设选题" in all_text  # intro from _create_group


async def test_export_admin_member_allowed(client: AsyncClient, class_setup, login) -> None:
    r = await client.get("/classes/me/groups/export", headers=await login(ADMIN_MEMBER))
    assert r.status_code == 200


async def test_export_member_forbidden(client: AsyncClient, class_setup, login) -> None:
    r = await client.get("/classes/me/groups/export", headers=await login(MEMBER1))
    assert r.status_code == 403


async def test_export_classless_forbidden(client: AsyncClient, class_setup, login) -> None:
    r = await client.get("/classes/me/groups/export", headers=await login("20240003001"))
    assert r.status_code == 403


# --- helpers ----------------------------------------------------------------


async def committee(login) -> dict[str, str]:
    return await login(COMMITTEE)
