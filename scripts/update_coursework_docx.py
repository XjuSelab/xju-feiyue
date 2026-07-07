from __future__ import annotations

import re
import shutil
import sys
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH


XJU_DOCX_SCRIPTS = Path.home() / ".codex" / "skills" / "xju-docx" / "scripts"
if XJU_DOCX_SCRIPTS.exists():
    sys.path.insert(0, str(XJU_DOCX_SCRIPTS))
    import xju_format
else:
    xju_format = None


ROOT = Path(__file__).resolve().parents[1]
DOCS = ROOT / "docs"
COURSEWORK = DOCS / "coursework"
TARGET = DOCS / "赵文彪-飞跃 .docx"

CHAPTERS = [
    COURSEWORK / "ch4-需求分析-笔记系统与社区互动.md",
    COURSEWORK / "ch5-概要设计-笔记系统与社区互动.md",
    COURSEWORK / "ch6-详细设计-笔记系统与社区互动.md",
    COURSEWORK / "ch7-系统测试-笔记系统与社区互动.md",
]


def remove_paragraph(paragraph) -> None:
    element = paragraph._element
    element.getparent().remove(element)


def remove_body_range(doc: Document, start_paragraph, end_paragraph) -> None:
    body = doc._body._element
    children = list(body.iterchildren())
    start_idx = children.index(start_paragraph._element)
    end_idx = children.index(end_paragraph._element)
    for element in children[start_idx:end_idx]:
        body.remove(element)


def insert_element_before(anchor, element) -> None:
    anchor._element.addprevious(element)


def add_before(doc: Document, anchor, text: str = "", style: str | None = None):
    paragraph = doc.add_paragraph(text)
    if style:
        paragraph.style = style
    apply_xju_paragraph_format(paragraph, style)
    insert_element_before(anchor, paragraph._element)
    return paragraph


def add_table_before(doc: Document, anchor, rows: list[list[str]]) -> None:
    if not rows:
        return
    width = max(len(row) for row in rows)
    normalized_rows = [
        [row[c_idx].strip() if c_idx < len(row) else "" for c_idx in range(width)]
        for row in rows
    ]
    if xju_format is not None:
        table = xju_format.three_line_table(
            doc,
            normalized_rows[0],
            normalized_rows[1:],
        )
    else:
        table = doc.add_table(rows=len(normalized_rows), cols=width)
        table.style = "Table Grid"
        for r_idx, row in enumerate(normalized_rows):
            for c_idx, value in enumerate(row):
                table.cell(r_idx, c_idx).text = value
    insert_element_before(anchor, table._element)
    add_before(doc, anchor, "")


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    i = start
    while i < len(lines) and lines[i].strip().startswith("|"):
        line = lines[i].strip()
        parts = [p.strip() for p in line.strip("|").split("|")]
        if not all(re.fullmatch(r":?-{3,}:?", p.replace(" ", "")) for p in parts):
            rows.append(parts)
        i += 1
    return rows, i


def clean_inline(text: str) -> str:
    text = text.replace("**", "")
    text = text.replace("`", "")
    text = text.replace("<br>", "\n")
    text = re.sub(r"<!--.*?-->", "", text)
    return text.strip()


def is_figure_line(text: str) -> bool:
    return "此处插入图" in text or re.match(r"^图\s*\d+[-.]\d+", text) is not None


def apply_xju_paragraph_format(paragraph, style: str | None = None) -> None:
    if xju_format is None:
        return

    if style == "图题":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in paragraph.runs:
            xju_format.set_run_font(
                run,
                xju_format.SIZE["五号"],
                bold=True,
            )
        xju_format.spacing_attrs(
            paragraph._p.get_or_add_pPr(),
            before_lines=0,
            after_lines=0,
            line=360,
        )
        return

    if style in {"Heading 1", "Heading 2", "Heading 3", "Heading 4"}:
        return

    for run in paragraph.runs:
        xju_format.set_run_font(run, xju_format.SIZE["小四"])
    ppr = paragraph._p.get_or_add_pPr()
    xju_format.spacing_attrs(ppr, line=360)
    if style not in {"List Bullet", "List Number"}:
        xju_format.first_line_chars(ppr, 200, 24)


def insert_markdown(doc: Document, anchor, md_text: str) -> None:
    lines = md_text.splitlines()
    i = 0
    while i < len(lines):
        raw = lines[i]
        line = raw.strip()
        if not line:
            add_before(doc, anchor, "")
            i += 1
            continue
        if line.startswith("|"):
            rows, i = parse_table(lines, i)
            add_table_before(doc, anchor, rows)
            continue
        if line.startswith("# "):
            add_before(doc, anchor, clean_inline(line[2:]), "Heading 1")
        elif line.startswith("## "):
            add_before(doc, anchor, clean_inline(line[3:]), "Heading 2")
        elif line.startswith("### "):
            add_before(doc, anchor, clean_inline(line[4:]), "Heading 3")
        elif line.startswith("#### "):
            add_before(doc, anchor, clean_inline(line[5:]), "Heading 4")
        elif is_figure_line(line):
            add_before(doc, anchor, clean_inline(line), "图题")
        elif line.startswith(">"):
            add_before(doc, anchor, clean_inline(line.lstrip("> ")), "Quote")
        elif line.startswith("- "):
            add_before(doc, anchor, clean_inline(line[2:]), "List Bullet")
        elif re.match(r"^\d+\.\s+", line):
            add_before(doc, anchor, clean_inline(re.sub(r"^\d+\.\s+", "", line)), "List Number")
        elif line.startswith("```"):
            block: list[str] = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                block.append(lines[i])
                i += 1
            add_before(doc, anchor, "\n".join(block))
        else:
            add_before(doc, anchor, clean_inline(line))
        i += 1


def looks_like_coursework_chapter_4(text: str) -> bool:
    return (
        "四" in text
        or "�ġ�" in text
        or "系统需求" in text
        or "ϵͳ�������" in text
    )


def looks_like_chapter_8(text: str) -> bool:
    return "八" in text or "�ˡ�" in text or "用户" in text


def remove_stray_pre_chapter4_block(doc: Document) -> None:
    body = doc._body._element
    children = list(body.iterchildren())
    paragraph_by_element = {paragraph._element: paragraph for paragraph in doc.paragraphs}

    tail_idx = None
    for idx, element in enumerate(children):
        paragraph = paragraph_by_element.get(element)
        if paragraph is not None and "专题计划包括四项" in paragraph.text:
            tail_idx = idx

    if tail_idx is None:
        return

    chapter4_idx = None
    for idx in range(tail_idx + 1, len(children)):
        paragraph = paragraph_by_element.get(children[idx])
        if (
            paragraph is not None
            and paragraph.style.name == "Heading 1"
            and looks_like_coursework_chapter_4(paragraph.text.strip())
        ):
            chapter4_idx = idx
            break

    if chapter4_idx is None or chapter4_idx <= tail_idx + 1:
        return

    stale_elements = children[tail_idx + 1 : chapter4_idx]
    stale_table_count = sum(
        1 for element in stale_elements if element.tag.rsplit("}", 1)[-1] == "tbl"
    )
    if stale_table_count < 5:
        return

    for element in stale_elements:
        body.remove(element)


def main() -> None:
    if not TARGET.exists():
        raise FileNotFoundError(TARGET)

    backup = TARGET.with_name(
        f"{TARGET.stem}.backup-{datetime.now().strftime('%Y%m%d-%H%M%S')}{TARGET.suffix}"
    )
    shutil.copy2(TARGET, backup)

    doc = Document(TARGET)
    if xju_format is not None:
        xju_format.setup_styles(doc)
        for section in doc.sections:
            xju_format.setup_section(section)
        xju_format.enable_update_fields(doc)
    remove_stray_pre_chapter4_block(doc)

    start = None
    end = None
    for idx, paragraph in enumerate(doc.paragraphs):
        text = paragraph.text.strip()
        if (
            start is None
            and idx > 240
            and paragraph.style.name == "Heading 1"
            and looks_like_coursework_chapter_4(text)
        ):
            start = idx
        if (
            start is not None
            and idx > start
            and paragraph.style.name == "Heading 1"
            and looks_like_chapter_8(text)
        ):
            end = idx
            break

    if start is None or end is None or start >= end:
        raise RuntimeError(f"Could not locate chapter 4-7 range: start={start}, end={end}")

    anchor = doc.paragraphs[end]
    remove_body_range(doc, doc.paragraphs[start], anchor)

    combined = "\n\n".join(path.read_text(encoding="utf-8") for path in CHAPTERS)
    insert_markdown(doc, anchor, combined)

    if xju_format is not None:
        xju_format.setup_styles(doc)
        for section in doc.sections:
            xju_format.setup_section(section)
        xju_format.enable_update_fields(doc)

    doc.save(TARGET)
    print(f"updated {TARGET}")
    print(f"backup {backup}")


if __name__ == "__main__":
    main()
