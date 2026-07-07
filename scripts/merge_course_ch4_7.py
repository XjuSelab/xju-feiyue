from __future__ import annotations

import copy
import shutil
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from PIL import Image, ImageDraw, ImageFont


ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "docs/course-design/赵文彪-飞跃.docx"
OVERVIEW = ROOT / "docs/course-design/feiyue-overview-software-engineering.png"

MODULES = [
    ("资料库与教务学分自查模块（丁水悦）", ROOT / "赵文彪-飞跃_丁水悦修改版.docx"),
    ("导师/院校库与 CCF 会议截稿模块（王艺婷）", ROOT / "schools-conferences-ch4-7-xju王艺婷 (1).docx"),
    ("用户、班级、小组、权限与任务管理模块（石建华）", ROOT / "飞跃合并版_班级管理_石建华_无代码.docx"),
]

CHAPTERS = [
    ("四、系统需求分析", ("四", "4"), ("五", "5")),
    ("五、系统概要设计", ("五", "5"), ("六", "6")),
    ("六、系统详细设计", ("六", "6"), ("七", "7")),
    ("七、系统测试", ("七", "7"), ("八", "8")),
]


def _font(size: int, bold: bool = False) -> ImageFont.FreeTypeFont:
    path = (
        "/usr/share/fonts/opentype/noto/NotoSansCJK-Bold.ttc"
        if bold
        else "/usr/share/fonts/opentype/noto/NotoSansCJK-Regular.ttc"
    )
    return ImageFont.truetype(path, size)


def _text_center(draw: ImageDraw.ImageDraw, box, text: str, font, fill="#1f2937") -> None:
    x1, y1, x2, y2 = box
    lines = text.split("\n")
    heights = []
    widths = []
    for line in lines:
        bb = draw.textbbox((0, 0), line, font=font)
        widths.append(bb[2] - bb[0])
        heights.append(bb[3] - bb[1])
    total_h = sum(heights) + (len(lines) - 1) * 8
    y = y1 + ((y2 - y1) - total_h) / 2
    for line, w, h in zip(lines, widths, heights):
        draw.text((x1 + ((x2 - x1) - w) / 2, y), line, font=font, fill=fill)
        y += h + 8


def _round_rect(draw, box, fill, outline="#374151", width=3, radius=20):
    draw.rounded_rectangle(box, radius=radius, fill=fill, outline=outline, width=width)


def _arrow(draw, start, end, fill="#374151", width=4):
    draw.line([start, end], fill=fill, width=width)
    x1, y1 = start
    x2, y2 = end
    # Simple arrowheads for horizontal and vertical connectors.
    if abs(x2 - x1) >= abs(y2 - y1):
        sign = 1 if x2 >= x1 else -1
        pts = [(x2, y2), (x2 - 16 * sign, y2 - 9), (x2 - 16 * sign, y2 + 9)]
    else:
        sign = 1 if y2 >= y1 else -1
        pts = [(x2, y2), (x2 - 9, y2 - 16 * sign), (x2 + 9, y2 - 16 * sign)]
    draw.polygon(pts, fill=fill)


def build_overview_image() -> None:
    img = Image.new("RGB", (1800, 1080), "#ffffff")
    d = ImageDraw.Draw(img)
    title = _font(44, True)
    head = _font(30, True)
    body = _font(25)
    small = _font(22)

    d.text((60, 38), "飞跃手册总览软工图", font=title, fill="#111827")
    d.text((60, 96), "需求、概要、详细设计与测试按业务域纵向切分，统一接入前后端分离平台", font=small, fill="#4b5563")

    users = (70, 195, 350, 430)
    frontend = (500, 160, 1300, 470)
    backend = (500, 560, 1300, 870)
    storage = (1430, 260, 1710, 760)
    external = (70, 610, 350, 860)

    _round_rect(d, users, "#eef2ff", "#4f46e5")
    _text_center(d, users, "用户角色\n学生 / 校友 / 访客\n管理员 / 超级管理员", body)

    _round_rect(d, frontend, "#ecfeff", "#0891b2")
    _text_center(
        d,
        (frontend[0], frontend[1] + 18, frontend[2], frontend[1] + 78),
        "React 单页应用",
        head,
    )
    features = [
        ("笔记系统\n（陶语涵待补）", (545, 265, 735, 345), "#fef3c7"),
        ("资料库 / 学分\n丁水悦", (760, 265, 950, 345), "#dcfce7"),
        ("导师 / 会议\n王艺婷", (975, 265, 1165, 345), "#dbeafe"),
        ("班级 / 权限\n石建华", (545, 365, 735, 445), "#fae8ff"),
        ("AI 辅助", (760, 365, 950, 445), "#ffedd5"),
        ("管理后台", (975, 365, 1165, 445), "#e5e7eb"),
    ]
    for text, box, fill in features:
        _round_rect(d, box, fill, "#6b7280", width=2, radius=14)
        _text_center(d, box, text, small)

    _round_rect(d, backend, "#f0fdf4", "#16a34a")
    _text_center(d, (backend[0], backend[1] + 18, backend[2], backend[1] + 78), "FastAPI 分层单体", head)
    layers = [
        ("routes\nHTTP 接口", (545, 665, 735, 760)),
        ("services\n业务规则", (805, 665, 995, 760)),
        ("db / models\n数据访问", (1065, 665, 1255, 760)),
    ]
    for text, box in layers:
        _round_rect(d, box, "#ffffff", "#65a30d", width=2, radius=14)
        _text_center(d, box, text, small)
    _arrow(d, (735, 712), (805, 712), "#65a30d")
    _arrow(d, (995, 712), (1065, 712), "#65a30d")

    _round_rect(d, storage, "#f8fafc", "#475569")
    _text_center(d, storage, "数据与文件\n主库 SQLite\n上传文件\nschools.sqlite\nconferences.sqlite\nmanifest / 备份", body)

    _round_rect(d, external, "#fff7ed", "#ea580c")
    _text_center(d, external, "外部系统\n教务成绩单\nccfddl / DDG\nDeepSeek\nsupervisor-claw\nHugging Face 备份", body)

    _arrow(d, (350, 310), (500, 310), "#374151")
    _arrow(d, (900, 470), (900, 560), "#374151")
    _arrow(d, (1300, 715), (1430, 510), "#374151")
    _arrow(d, (350, 735), (500, 735), "#ea580c")
    _arrow(d, (1300, 310), (1430, 390), "#475569")

    d.text((560, 505), "HTTPS / JSON API，经 nginx 反向代理", font=small, fill="#4b5563")
    d.text((1320, 640), "只读快照与主业务库隔离", font=small, fill="#4b5563")
    d.text((390, 768), "授权导入 / 爬取 / AI 调用 / 灾备", font=small, fill="#4b5563")

    OVERVIEW.parent.mkdir(parents=True, exist_ok=True)
    img.save(OVERVIEW)


def heading_text(p) -> str:
    return "".join(p.text.split())


def is_heading(p, starts: tuple[str, ...], level: int | None = None) -> bool:
    text = heading_text(p)
    if not p.style or not p.style.name.startswith("Heading"):
        return False
    if level is not None and p.style.name != f"Heading {level}":
        return False
    return any(text.startswith(s) for s in starts)


def section_elements(doc: Document, start_keys: tuple[str, ...], end_keys: tuple[str, ...]):
    body = doc._body._element
    children = list(body)
    start = end = None
    for i, child in enumerate(children):
        if child.tag != qn("w:p"):
            continue
        p = doc.paragraphs[[x._element for x in doc.paragraphs].index(child)]
        if start is None and is_heading(p, start_keys, level=1):
            start = i
            continue
        if start is not None and is_heading(p, end_keys, level=1):
            end = i
            break
    if start is None:
        raise RuntimeError(f"未找到章节起点 {start_keys}")
    if end is None:
        end = len(children)
    return children[start + 1 : end]


def remap_image_relationships(src_doc: Document, dst_doc: Document, element) -> None:
    for node in element.xpath(".//*[@r:embed]"):
        old = node.get(qn("r:embed"))
        part = src_doc.part.related_parts.get(old)
        if part is None or not hasattr(part, "blob"):
            continue
        new_rid, _ = dst_doc.part.get_or_add_image(BytesIO(part.blob))
        node.set(qn("r:embed"), new_rid)


def demote_heading_styles(element) -> None:
    for p in element.xpath(".//w:p"):
        pstyle = p.find(qn("w:pPr") + "/" + qn("w:pStyle"))
        if pstyle is None:
            continue
        val = pstyle.get(qn("w:val"))
        if val == "Heading2":
            pstyle.set(qn("w:val"), "Heading3")
        elif val == "Heading3":
            pstyle.set(qn("w:val"), "Heading4")


def insert_existing_before(marker, element) -> None:
    marker.addprevious(element)


def add_para_before(doc: Document, marker, text: str = "", style: str | None = None, align=None):
    p = doc.add_paragraph(text, style=style)
    if align is not None:
        p.alignment = align
    el = p._element
    doc._body._element.remove(el)
    marker.addprevious(el)
    return el


def add_picture_before(doc: Document, marker, path: Path, width_inches: float = 6.2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run()
    run.add_picture(str(path), width=Inches(width_inches))
    el = p._element
    doc._body._element.remove(el)
    marker.addprevious(el)
    return el


def remove_target_chapters(doc: Document):
    body = doc._body._element
    children = list(body)
    para_by_el = {p._element: p for p in doc.paragraphs}
    start = end = None
    for i, child in enumerate(children):
        p = para_by_el.get(child)
        if p is None:
            continue
        if start is None and is_heading(p, ("四", "4"), level=1):
            start = i
            continue
        if start is not None and is_heading(p, ("八", "8"), level=1):
            end = i
            break
    if start is None or end is None:
        raise RuntimeError("未找到总文档第 4-7 章或第 8 章边界")
    marker = children[end]
    for child in children[start:end]:
        body.remove(child)
    return marker


def add_caption_before(doc: Document, marker, text: str):
    el = add_para_before(doc, marker, text, style="图题", align=WD_ALIGN_PARAGRAPH.CENTER)
    return el


def add_tao_placeholder(doc: Document, marker, chapter_title: str):
    add_para_before(
        doc,
        marker,
        "陶语涵负责的笔记系统模块（待补）",
        style="Heading 2",
    )
    add_para_before(
        doc,
        marker,
        f"【陶语涵待补】本节对应{chapter_title}中笔记系统（写作、展示、评论、点赞）的内容。"
        "按照第 3 章分工，后续需补充 /write 与 /browse 相关的用例、图表、设计说明、核心代码和测试记录；"
        "当前合并稿仅保留空位，不代写未提交材料。",
    )


def merge() -> None:
    build_overview_image()
    timestamp = datetime.now().strftime("%Y%m%d-%H%M%S")
    backup = TARGET.with_name(f"{TARGET.stem}.bak-{timestamp}{TARGET.suffix}")
    shutil.copy2(TARGET, backup)

    dst = Document(str(TARGET))
    marker = remove_target_chapters(dst)

    source_docs = [(name, Document(str(path))) for name, path in MODULES]
    for chapter_title, start_keys, end_keys in CHAPTERS:
        add_para_before(dst, marker, chapter_title, style="Heading 1", align=WD_ALIGN_PARAGRAPH.CENTER)

        if chapter_title.startswith("四"):
            add_para_before(
                dst,
                marker,
                "本章在前三章项目背景、可行性与开发计划的基础上，按业务域汇总第 4 至第 7 章的模块材料。"
                "当前已完成资料库与教务学分自查、导师/院校库与 CCF 会议截稿、用户班级小组权限与任务管理三个模块；"
                "笔记系统由陶语涵负责，因尚未提交材料，本合并稿保留空位。",
            )
            add_para_before(
                dst,
                marker,
                "飞跃手册总览软工图如图4-1所示。该图把系统从用户角色、前端功能域、后端分层单体、数据与文件存储、外部系统五个层面展开，"
                "用于说明各成员模块如何在统一平台中协同运行。",
            )
            add_picture_before(dst, marker, OVERVIEW)
            add_caption_before(dst, marker, "图4-1　飞跃手册总览软工图")
            add_para_before(
                dst,
                marker,
                "图4-1中，左侧为系统使用者与外部系统，中心为 React 单页应用和 FastAPI 分层单体，右侧为主库、上传文件、只读快照库和备份文件。"
                "丁水悦模块主要位于资料库、学分自查和上传文件链路；王艺婷模块主要位于导师、会议与只读快照库链路；"
                "石建华模块主要位于鉴权、班级、小组、权限与任务链路；陶语涵模块对应笔记写作、浏览、评论和点赞，当前保留待补。"
                "图中箭头表示 HTTPS 请求、服务层调用、只读查询、外部数据导入和灾备流向。",
            )

        add_tao_placeholder(dst, marker, chapter_title)

        module_index = 2
        for module_name, src_doc in source_docs:
            add_para_before(dst, marker, f"{module_index}．{module_name}", style="Heading 2")
            for src_el in section_elements(src_doc, start_keys, end_keys):
                el = copy.deepcopy(src_el)
                remap_image_relationships(src_doc, dst, el)
                demote_heading_styles(el)
                insert_existing_before(marker, el)
            module_index += 1

    for p in dst.paragraphs:
        if p.style and p.style.name == "Normal":
            p.paragraph_format.line_spacing = 1.5
            for run in p.runs:
                run.font.size = Pt(12)

    dst.save(str(TARGET))
    print(f"merged: {TARGET}")
    print(f"backup: {backup}")
    print(f"overview: {OVERVIEW}")


if __name__ == "__main__":
    merge()
