from __future__ import annotations

from pathlib import Path
import re
import shutil
from datetime import datetime

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt


ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "docs/course-design/赵文彪-飞跃.docx"

BODY_FONT = "宋体"
LATIN_FONT = "Times New Roman"
CODE_FONT = "Consolas"


CODE_SUMMARIES = {
    "MAX_UPLOAD_BYTES": "关键逻辑：限制单文件 50MB；按分块读取上传内容；校验扩展名、MIME 与文件魔数；写入临时文件后原子移动；异常时清理临时文件，避免孤儿文件。",
    "_TTL_SECONDS": "关键逻辑：成绩单 PDF 仅保存在内存字典中；设置 5 分钟 TTL 与 10MB 上限；上传后返回一次性 token；读取时使用 pop 取后即删，并清理过期暂存件。",
    'if position == "inside"': "关键逻辑：文件树重排前校验目标是否为文件夹；沿 parent_id 追溯祖先链，禁止把节点拖入自身或后代；通过后更新 parent_id 与 sort_order。",
    "export async function loadTextItems": "关键逻辑：浏览器端读取 PDF ArrayBuffer；使用 pdf.js 逐页提取 text items；按坐标恢复行列；后续由 parseTranscript 识别课程、成绩与通识模块。",
}


def set_run_font(run, east=BODY_FONT, latin=LATIN_FONT, size=None):
    run.font.name = latin
    if size is not None:
        run.font.size = size
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), east)
    rfonts.set(qn("w:ascii"), latin)
    rfonts.set(qn("w:hAnsi"), latin)
    rfonts.set(qn("w:cs"), latin)


def normalize_fonts(doc: Document) -> None:
    for style_name, size in [
        ("Normal", Pt(12)),
        ("Heading 1", Pt(16)),
        ("Heading 2", Pt(15)),
        ("Heading 3", Pt(14)),
        ("Heading 4", Pt(12)),
        ("图题", Pt(10.5)),
        ("表题", Pt(10.5)),
    ]:
        if style_name in doc.styles:
            style = doc.styles[style_name]
            style.font.name = LATIN_FONT
            style.font.size = size
            style._element.rPr.rFonts.set(qn("w:eastAsia"), BODY_FONT)
            style._element.rPr.rFonts.set(qn("w:ascii"), LATIN_FONT)
            style._element.rPr.rFonts.set(qn("w:hAnsi"), LATIN_FONT)

    for para in doc.paragraphs:
        size = None
        if para.style.name == "Heading 1":
            size = Pt(16)
        elif para.style.name == "Heading 2":
            size = Pt(15)
        elif para.style.name == "Heading 3":
            size = Pt(14)
        elif para.style.name in {"图题", "表题"}:
            size = Pt(10.5)
        for run in para.runs:
            set_run_font(run, size=size)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    for run in para.runs:
                        set_run_font(run, size=Pt(10.5))


def clean_text(text: str) -> str:
    return re.sub(r"\s+", " ", text.replace("\u3000", " ")).strip()


def shorten(text: str, limit: int) -> str:
    text = clean_text(text)
    if len(text) <= limit:
        return text

    parts = re.split(r"(?<=[。；;])", text)
    out = ""
    for part in parts:
        if not part:
            continue
        if len(out) + len(part) <= limit:
            out += part
        else:
            break
    if len(out) < max(18, limit // 3):
        out = text[: limit - 1]
    out = out.rstrip("，,；;。 ")
    return out + "。"


def set_cell_text(cell, text: str, size=Pt(10.5), code=False) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    if code:
        set_run_font(run, east=BODY_FONT, latin=CODE_FONT, size=Pt(9))
    else:
        set_run_font(run, size=size)


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, cm: float) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(int(cm * 567)))


def set_table_width(table, col_widths: list[float]) -> None:
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            if idx < len(col_widths):
                set_cell_width(cell, col_widths[idx])


def compact_tables(doc: Document) -> None:
    for ti, table in enumerate(doc.tables):
        rows = len(table.rows)
        cols = len(table.columns) if rows else 0
        if not rows or not cols:
            continue

        # Header formatting.
        for cell in table.rows[0].cells:
            set_cell_shading(cell, "F2F2F2")
            for para in cell.paragraphs:
                para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                for run in para.runs:
                    run.bold = True
                    set_run_font(run, size=Pt(10.5))

        # Replace oversized code listing tables with concise implementation summaries.
        if rows == 1 and cols == 1:
            raw = table.cell(0, 0).text
            for marker, summary in CODE_SUMMARIES.items():
                if marker in raw:
                    set_cell_text(table.cell(0, 0), summary, size=Pt(10.5))
                    break

        # Compress long cells so narrow columns remain readable.
        if cols >= 5:
            limit = 62
        elif cols == 4:
            limit = 78
        elif cols == 3:
            limit = 95
        elif cols == 2:
            limit = 135
        else:
            limit = 180

        for ri, row in enumerate(table.rows):
            for ci, cell in enumerate(row.cells):
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                text = clean_text(cell.text)
                if ri == 0 or not text:
                    continue
                if len(text) > limit:
                    set_cell_text(cell, shorten(text, limit))

        if cols == 2:
            set_table_width(table, [3.2, 12.2])
        elif cols == 3:
            set_table_width(table, [3.0, 5.0, 7.4])
        elif cols == 4:
            set_table_width(table, [2.5, 4.0, 4.3, 4.6])
        elif cols == 5:
            set_table_width(table, [2.2, 3.0, 3.4, 3.6, 3.2])
        elif cols == 6:
            set_table_width(table, [2.0, 3.0, 2.5, 2.7, 2.8, 2.4])

        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    para.paragraph_format.space_before = Pt(0)
                    para.paragraph_format.space_after = Pt(0)
                    para.paragraph_format.line_spacing = 1.15
                    for run in para.runs:
                        if cols >= 4 or rows > 8:
                            run.font.size = Pt(9.5)
                        else:
                            run.font.size = Pt(10.5)
                        # Preserve explicit code summary size.
                        if rows == 1 and cols == 1:
                            run.font.size = Pt(10)


def main() -> None:
    backup = TARGET.with_name(f"{TARGET.stem}.tables-bak-{datetime.now():%Y%m%d-%H%M%S}{TARGET.suffix}")
    shutil.copy2(TARGET, backup)
    doc = Document(str(TARGET))
    normalize_fonts(doc)
    compact_tables(doc)
    doc.save(str(TARGET))
    print(f"tuned: {TARGET}")
    print(f"backup: {backup}")


if __name__ == "__main__":
    main()
