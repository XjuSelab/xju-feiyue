from __future__ import annotations

import copy
import shutil
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt


ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "docs/course-design/赵文彪-飞跃.docx"
SOURCE = ROOT / "赵文彪-飞跃-陶语涵.docx"

BODY_FONT = "宋体"
LATIN_FONT = "Times New Roman"


def set_run_font(run, size=Pt(12), bold=False):
    run.font.name = LATIN_FONT
    run.font.size = size
    run.bold = bold
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:eastAsia"), BODY_FONT)
    rfonts.set(qn("w:ascii"), LATIN_FONT)
    rfonts.set(qn("w:hAnsi"), LATIN_FONT)
    rfonts.set(qn("w:cs"), LATIN_FONT)


def para_map(doc):
    return {p._element: p for p in doc.paragraphs}


def heading_text(p):
    return "".join(p.text.split())


def add_para_before(doc, marker, text, style="Normal", align=None):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    run = p.add_run(text)
    if style == "Heading 2":
        set_run_font(run, Pt(15), True)
    elif style == "Heading 3":
        set_run_font(run, Pt(14), True)
    elif style in {"表题", "图题"}:
        set_run_font(run, Pt(10.5), True)
    else:
        set_run_font(run, Pt(12))
    el = p._element
    doc._body._element.remove(el)
    marker.addprevious(el)
    return el


def add_table_before(doc, marker, caption, headers, rows):
    add_para_before(doc, marker, caption, style="表题", align=WD_ALIGN_PARAGRAPH.CENTER)
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        set_run_font(run, Pt(10.5), True)
    for row in rows:
        cells = table.add_row().cells
        for i, text in enumerate(row):
            cells[i].text = ""
            run = cells[i].paragraphs[0].add_run(text)
            set_run_font(run, Pt(10.5))
    el = table._element
    doc._body._element.remove(el)
    marker.addprevious(el)
    return el


def remap_image_relationships(src_doc: Document, dst_doc: Document, element) -> None:
    for node in element.xpath(".//*[@r:embed]"):
        old = node.get(qn("r:embed"))
        part = src_doc.part.related_parts.get(old)
        if part is None or not hasattr(part, "blob"):
            continue
        new_rid, _ = dst_doc.part.get_or_add_image(BytesIO(part.blob))
        node.set(qn("r:embed"), new_rid)


def insert_source_figure(dst_doc: Document, src_doc: Document, marker, caption_contains: str) -> None:
    for i, para in enumerate(src_doc.paragraphs):
        if caption_contains not in para.text:
            continue
        image_para = None
        for j in range(i - 1, max(-1, i - 8), -1):
            if "<w:drawing" in src_doc.paragraphs[j]._element.xml or "<w:pict" in src_doc.paragraphs[j]._element.xml:
                image_para = src_doc.paragraphs[j]
                break
        if image_para is None:
            return
        img_el = copy.deepcopy(image_para._element)
        remap_image_relationships(src_doc, dst_doc, img_el)
        marker.addprevious(img_el)
        add_para_before(dst_doc, marker, para.text.strip(), style="图题", align=WD_ALIGN_PARAGRAPH.CENTER)
        return


def find_tao_range(doc, chapter_start):
    children = list(doc._body._element)
    pmap = para_map(doc)
    in_chapter = False
    start = end = None
    for i, child in enumerate(children):
        p = pmap.get(child)
        if p is None:
            continue
        if not in_chapter and p.style.name == "Heading 1" and heading_text(p).startswith(chapter_start):
            in_chapter = True
            continue
        if not in_chapter:
            continue
        if p.style.name == "Heading 1":
            break
        if start is None and p.style.name == "Heading 2" and "陶语涵" in p.text:
            start = i
            continue
        if start is not None and p.style.name in {"Heading 1", "Heading 2"}:
            end = i
            break
    if start is None or end is None:
        raise RuntimeError(f"cannot find Tao range for chapter {chapter_start}")
    return start, end


def replace_range(doc, chapter_start, writer):
    body = doc._body._element
    children = list(body)
    start, end = find_tao_range(doc, chapter_start)
    marker = children[end]
    for child in children[start:end]:
        body.remove(child)
    writer(doc, marker)


def write_ch4(doc, marker):
    src = Document(str(SOURCE))
    add_para_before(doc, marker, "1．笔记系统与社区互动模块（陶语涵）", "Heading 2")
    add_para_before(
        doc,
        marker,
        "本模块负责飞跃手册的内容生产与社区互动入口，覆盖 /write、/browse、/notes、/collections 与用户经验相关接口。"
        "需求重点不是单独追求功能数量，而是保证“写作、发布、浏览、互动、治理”形成闭环，并与平台统一鉴权、AI 降级和主库模型保持一致。",
    )
    add_table_before(
        doc,
        marker,
        "表4-1　笔记系统与社区互动模块核心需求",
        ["需求域", "主要功能", "验收要点"],
        [
            ("笔记写作", "Markdown 草稿、自动保存、发布、AI 选段润色和摘要兜底。", "草稿可恢复；发布前校验标题、正文和分类；AI 失败不阻断发布。"),
            ("笔记浏览", "信息流、分类筛选、关键词检索、详情页和锚点定位。", "匿名访客可读公开内容；分页稳定；详情页能展示摘要、正文和互动计数。"),
            ("社区互动", "评论、楼中楼回复、点赞/点踩互斥、收藏、合集上下文。", "写操作要求登录；重复表态幂等；单篇笔记至多属于一个合集。"),
            ("成长与治理", "每日签到、经验流水、举报与 AI 审查规划。", "签到每日一次；经验记录可追踪；举报和人工裁决作为后续治理扩展。"),
        ],
    )
    add_para_before(
        doc,
        marker,
        "本模块的主要参与者包括匿名访客、注册用户、笔记作者和管理员。匿名访客只能浏览公开内容；注册用户可以评论、收藏、表态和签到；笔记作者可以维护本人草稿、已发布笔记和合集；管理员负责后续举报处理、内容隐藏和异常数据维护。",
    )
    add_para_before(
        doc,
        marker,
        "从业务流程看，笔记系统首先解决“经验如何被写下来”的问题，再解决“经验如何被找到和讨论”的问题。作者在写作页创建草稿，系统在编辑过程中持续自动保存；发布时后端统一校验标题、正文、分类和作者身份，并尝试生成摘要。读者在浏览页通过分类、关键词、标签和排序方式发现内容，在详情页完成阅读、评论、收藏和表态。该流程使内容生产、内容消费和社区反馈形成闭环。",
    )
    add_para_before(
        doc,
        marker,
        "非功能需求主要体现在三方面。第一，可靠性：草稿自动保存和 AI 摘要降级必须保证用户不会因网络波动或模型失败丢失编辑成果。第二，权限一致性：所有写操作都必须以后端当前用户为准，前端隐藏按钮只能作为体验优化，不能作为安全边界。第三，可维护性：评论、表态、收藏、合集和经验系统均应通过独立服务函数实现，避免页面组件直接拼接复杂业务规则。",
    )
    insert_source_figure(doc, src, marker, "图 4-1 笔记系统与社区互动模块系统流程图")
    add_table_before(
        doc,
        marker,
        "表4-2　笔记系统与社区互动模块数据字典摘要",
        ["数据对象", "关键字段", "说明"],
        [
            ("Note", "id、title、content、summary、category、author_sid、status", "已发布笔记主体，承载浏览、检索、互动和详情展示。"),
            ("Draft", "id、author_sid、title、content、autosaved_at", "写作页草稿，支持自动保存和发布前校验。"),
            ("Comment", "id、note_id、parent_id、author_sid、anchor、content", "评论与楼中楼回复，parent_id 限制为两层结构。"),
            ("Interaction", "note_id、user_sid、type", "点赞、点踩、收藏等用户行为，依赖唯一约束保证幂等。"),
            ("Collection", "id、owner_sid、title、description", "笔记合集；条目表限定单篇笔记最多归属一个合集。"),
        ],
    )
    insert_source_figure(doc, src, marker, "图 4-2 顶层数据流图")
    add_table_before(
        doc,
        marker,
        "表4-3　核心用例摘要",
        ["用例", "触发条件", "主成功场景", "异常场景"],
        [
            ("发布笔记", "作者完成草稿编辑后点击发布。", "系统校验草稿并写入公开笔记，返回详情页。", "标题缺失、正文为空或作者不一致时拒绝发布。"),
            ("发表评论", "登录用户在详情页提交评论。", "系统保存顶层评论或楼内回复，并刷新评论列表。", "未登录、笔记不可见或三级嵌套时返回错误。"),
            ("收藏笔记", "登录用户点击收藏按钮。", "系统写入收藏记录并更新收藏状态。", "重复收藏保持幂等，不产生重复数据。"),
            ("合集导航", "用户打开属于合集的笔记详情。", "系统返回合集名称、前后篇和排序信息。", "笔记不在合集时返回空上下文。"),
        ],
    )
    insert_source_figure(doc, src, marker, "图 4-5 笔记系统与社区互动模块用例图")
    add_para_before(
        doc,
        marker,
        "需求分析结论：笔记系统是平台内容沉淀的核心模块，必须优先保证发布链路可靠、浏览链路清晰、互动权限一致。AI 能力只作为辅助增强，所有依赖外部模型的流程均需具备本地兜底或失败降级。",
    )
    add_table_before(
        doc,
        marker,
        "表4-4　需求优先级与边界",
        ["优先级", "需求范围", "边界说明"],
        [
            ("高", "草稿、发布、浏览、详情、评论、点赞和收藏。", "属于本模块课程验收的核心闭环，必须可演示、可测试。"),
            ("中", "合集、签到、经验流水和个人中心展示。", "用于增强用户留存，后端能力优先于前端完整入口。"),
            ("中", "AI 润色、AI 摘要和摘要流式输出。", "作为写作辅助，不允许因为上游失败导致发布失败。"),
            ("低", "举报、AI 审查、拉黑、等级徽章和治理后台。", "保留数据模型和设计口径，作为后续迭代，不计为当前已完成项。"),
        ],
    )
    add_para_before(
        doc,
        marker,
        "数据流方面，本模块可以划分为三条主线。第一条是写作主线，草稿数据从前端编辑器进入 Draft，再经发布服务转化为 Note。第二条是互动主线，用户对 Note 产生评论、表态、收藏和合集归属，这些行为既影响详情页展示，也影响后续个人中心统计。第三条是辅助主线，AI 服务为写作和摘要提供增强能力，签到和经验系统为用户持续使用提供反馈。三条主线都围绕 Note 展开，因此需求分析中把笔记实体作为核心数据对象。",
    )
    add_para_before(
        doc,
        marker,
        "在需求取舍上，本模块将“可持续沉淀经验”作为第一目标，而不是把社区功能做成复杂论坛。课程经验、保研记录、工具教程和竞赛复盘都属于长文本内容，最需要稳定编辑、分类检索和后续讨论；因此评论、收藏和合集围绕笔记展开，不单独设计帖子、话题和圈子等额外对象。这样可以降低模型复杂度，也便于和现有 Markdown 内容导入流程对接。",
    )
    add_para_before(
        doc,
        marker,
        "对用户体验而言，写作页要尽量减少打断：自动保存应在后台完成，AI 润色只处理用户选中的片段，摘要生成可以在发布时自动触发但不能阻塞主流程。浏览页则强调快速判断内容价值，卡片中应展示标题、分类、摘要、作者和互动计数。详情页承担最长停留时间，评论区、收藏按钮和合集导航的位置应保持稳定，避免用户阅读过程中频繁跳动。",
    )


def write_ch5(doc, marker):
    src = Document(str(SOURCE))
    add_para_before(doc, marker, "1．笔记系统与社区互动模块（陶语涵）", "Heading 2")
    add_para_before(
        doc,
        marker,
        "概要设计采用前后端分离和分层单体结构。前端按写作、浏览、详情、评论、合集和个人中心划分 feature；后端按 routes、services、schemas、models 分层，所有数据库写入均经服务层校验权限和状态。",
    )
    add_para_before(
        doc,
        marker,
        "模块内部以 Note 为中心组织对象关系。Draft 面向写作过程，Note 面向发布后的公开内容，Comment 和 Interaction 面向读者反馈，Collection 面向作者对系列内容的组织，XPEvent 面向用户成长记录。这些对象都通过用户主键与鉴权模块关联，因此接口设计必须避免信任客户端传入的 author_sid，而应统一使用服务端解析出的当前用户。",
    )
    insert_source_figure(doc, src, marker, "图 5-1 笔记系统与社区互动模块总体架构")
    add_table_before(
        doc,
        marker,
        "表5-1　功能需求与程序结构映射",
        ["功能域", "前端结构", "后端结构", "设计重点"],
        [
            ("写作发布", "WritePage、DraftEditor、AIComposePanel", "routes/notes、services/notes、Draft/Note", "草稿自动保存、发布校验、摘要兜底。"),
            ("浏览详情", "BrowsePage、NoteDetail、MarkdownRenderer", "GET /notes、GET /notes/{id}", "公开读、分页、分类筛选和详情装配。"),
            ("评论互动", "CommentSection、InteractionButtons", "routes/comments、routes/interactions", "两层回复、赞踩互斥、收藏幂等。"),
            ("合集经验", "CollectionPanel、Profile widgets", "routes/collections、routes/auth", "合集单归属、签到幂等、经验流水。"),
        ],
    )
    add_table_before(
        doc,
        marker,
        "表5-2　数据库与约束设计摘要",
        ["表名", "关系", "关键约束"],
        [
            ("notes", "User 1:N Note", "作者外键；状态字段控制 visible/deleted；分类和标题非空。"),
            ("drafts", "User 1:N Draft", "草稿归作者所有；发布后可删除或转为历史状态。"),
            ("comments", "Note 1:N Comment，自引用 parent_id", "限制两层楼中楼；删除顶层评论时级联清理回复。"),
            ("note_interactions", "User N:M Note", "同一用户对同一笔记的同类操作唯一；赞踩互斥由服务层处理。"),
            ("collections / collection_entries", "User 1:N Collection，Collection 1:N Entry", "entry.note_id 唯一，保证单篇笔记只属于一个合集。"),
        ],
    )
    add_table_before(
        doc,
        marker,
        "表5-3　运行控制与补救措施",
        ["场景", "控制策略", "补救方式"],
        [
            ("草稿保存失败", "前端保留编辑器当前内容并提示重试。", "用户可重新触发保存；后端日志记录失败原因。"),
            ("AI 摘要失败", "摘要生成不作为发布事务的必要条件。", "使用正文前若干字符或空摘要兜底，发布主流程继续。"),
            ("并发表态", "服务层在同一用户与笔记维度执行互斥更新。", "最终以数据库记录为准，前端重新拉取计数。"),
            ("合集冲突", "collection_entries.note_id 设置唯一约束。", "返回明确冲突提示，引导作者先移出原合集。"),
        ],
    )
    add_para_before(
        doc,
        marker,
        "出错处理方面，发布失败、AI 超时、评论越权、重复签到和合集归属冲突均返回明确业务错误；前端保持原输入状态，允许用户修改后重试。对 AI 润色和摘要生成，系统优先使用流式能力，失败时回退普通请求，再失败时使用本地摘要占位，避免用户流程被外部服务阻塞。",
    )
    add_para_before(
        doc,
        marker,
        "与其他模块的接口边界也需要明确。资料库、导师库和班级空间可以在页面层链接到笔记内容，但不直接写入 notes 表；用户头像、昵称和角色来自鉴权模块；AI 服务仅通过统一封装调用，避免在写作页、发布服务和摘要接口中重复维护密钥、超时和降级逻辑。",
    )
    add_table_before(
        doc,
        marker,
        "表5-4　模块接口边界",
        ["相邻模块", "交互内容", "边界约束"],
        [
            ("鉴权模块", "当前用户、角色、昵称、头像、JWT 校验。", "笔记模块只消费当前用户身份，不维护密码和角色。"),
            ("AI 辅助模块", "润色、摘要、审查分类等模型调用。", "统一超时、限流和降级，业务层只处理结果。"),
            ("资料库模块", "资料详情可链接相关笔记或在笔记中引用资料。", "资料文件权限仍由资料库 owner 规则控制。"),
            ("管理后台", "后续内容隐藏、举报裁决和异常数据处理。", "普通用户接口不得暴露管理员处理能力。"),
        ],
    )
    add_para_before(
        doc,
        marker,
        "概要设计阶段还需要控制文档表达粒度。本模块功能较多，但课程文档应突出可运行闭环：写作发布、阅读互动和数据一致性。对于尚未完整上线的治理能力，文档只说明预留字段、状态流转和扩展方向，避免把规划项写成已交付功能。",
    )
    add_para_before(
        doc,
        marker,
        "运行时组合上，浏览类接口以只读查询为主，优先保证响应稳定和分页一致；写作类接口以事务写入为主，优先保证草稿、笔记和作者关系一致；互动类接口以小事务和唯一约束为主，优先保证幂等与并发安全。这样划分后，前端可以根据接口类型采用不同缓存策略：浏览列表可短时间缓存，发布和互动写操作完成后立即使相关查询失效并重新拉取。",
    )
    add_para_before(
        doc,
        marker,
        "数据库设计没有拆成多个独立库，而是沿用平台主库。原因是笔记、评论、收藏、合集和经验都与 User 强相关，跨库会增加事务一致性和查询装配成本。与导师库、会议库这种外部只读快照不同，笔记系统是高频写入模块，放在主业务库中更便于统一迁移、备份和权限审计。",
    )
    add_para_before(
        doc,
        marker,
        "缓存与刷新策略也需要纳入概要设计。信息流列表可以按分类、关键词和排序条件缓存；详情页在发表评论、点赞或收藏后应局部刷新计数与当前用户状态；草稿页不应依赖列表缓存，而应以草稿详情接口返回的最新内容为准。前端通过 TanStack Query 管理服务端状态，可以在写操作成功后精确失效相关 query，避免全局刷新造成页面闪烁。",
    )
    add_para_before(
        doc,
        marker,
        "权限设计遵循最小授权原则。匿名访客只拥有读公开笔记的权限；注册用户可以对公开笔记产生评论、收藏和表态；笔记作者可以编辑或删除本人内容；管理员只在治理功能中介入。任何涉及作者身份的接口都不接受客户端传入的 author_sid 作为可信依据，而是使用 JWT 解析出的当前用户，避免用户伪造请求越权修改他人内容。",
    )


def write_ch6(doc, marker):
    src = Document(str(SOURCE))
    add_para_before(doc, marker, "1．笔记系统与社区互动模块（陶语涵）", "Heading 2")
    add_para_before(
        doc,
        marker,
        "详细设计围绕四条关键链路展开：草稿发布、评论创建、赞踩收藏、合集上下文查询。各链路都遵循“前端校验只提升体验，后端校验决定权限和数据一致性”的原则。",
    )
    add_para_before(
        doc,
        marker,
        "草稿发布链路中，前端负责即时校验和编辑体验，后端负责最终一致性。后端接收发布请求后先读取草稿并确认作者身份，再执行字段校验；如果摘要为空，则调用 AI 摘要服务，服务失败时使用降级摘要。最后在一个数据库事务中创建 Note 并处理草稿状态，避免出现已公开但找不到作者或分类的异常记录。",
    )
    add_para_before(
        doc,
        marker,
        "评论链路采用平铺存储、逻辑分层展示的方式。数据库只保存 parent_id 和 reply_to 等必要字段，服务层限制 parent_id 只能指向顶层评论；前端根据 parent_id 将评论组装为顶层评论和楼内回复。这样既满足课程文档中“楼中楼”需求，又避免无限递归导致分页、删除和权限校验复杂化。",
    )
    insert_source_figure(doc, src, marker, "图 6-1 程序结构图")
    add_table_before(
        doc,
        marker,
        "表6-1　核心处理流程",
        ["流程", "主要步骤", "一致性控制"],
        [
            ("草稿发布", "读取草稿→校验标题/正文/分类→生成或兜底摘要→写入 Note→返回详情。", "发布操作只允许草稿作者执行；AI 失败不回滚发布主流程。"),
            ("创建评论", "校验登录→检查笔记可见→校验 parent_id 层级→写入 Comment。", "parent_id 只能指向顶层评论，避免无限嵌套。"),
            ("赞踩收藏", "接收 type→检查目标笔记→更新 interaction→刷新计数。", "点赞和点踩互斥；收藏独立幂等。"),
            ("合集上下文", "按 note_id 查询 entry→装配 collection 与前后篇→返回导航信息。", "只有合集创建者能维护条目；普通访客只能读取公开上下文。"),
        ],
    )
    add_table_before(
        doc,
        marker,
        "表6-2　关键接口设计",
        ["接口", "输入", "输出"],
        [
            ("POST /notes/drafts/{id}/publish", "草稿 id、标题、正文、分类、标签", "发布后的 Note 详情或校验错误。"),
            ("POST /notes/{id}/comments", "评论正文、anchor、parent_id", "评论记录及楼层关系。"),
            ("POST /notes/{id}/interaction", "like、dislike、favorite 等操作类型", "最新交互状态和计数。"),
            ("GET /notes/{id}/collection", "笔记 id", "所在合集、前一篇、后一篇和排序信息。"),
            ("POST /auth/checkin", "当前登录用户", "签到结果、经验变化和当日状态。"),
        ],
    )
    insert_source_figure(doc, src, marker, "图 6-2 写作页界面")
    add_table_before(
        doc,
        marker,
        "表6-3　核心函数职责",
        ["函数/组件", "所在层", "职责"],
        [
            ("publish_draft", "services/notes", "校验草稿归属、生成摘要、创建公开笔记。"),
            ("create_comment", "services/comments", "校验笔记可见性、楼层深度和评论作者。"),
            ("toggle_interaction", "routes/interactions", "处理点赞、点踩、收藏的幂等与互斥规则。"),
            ("get_collection_context", "routes/collections", "根据 note_id 返回合集与前后篇导航。"),
            ("CommentSection", "frontend/features/notes", "展示评论列表、提交评论并处理局部刷新。"),
        ],
    )
    insert_source_figure(doc, src, marker, "图 6-3 信息流浏览页界面")
    add_para_before(
        doc,
        marker,
        "实现结果上，草稿发布、浏览、评论、赞踩、收藏、合集和签到接口已有代码或测试依据；举报、AI 审查、拉黑与等级徽章属于治理扩展，文档中保留设计口径，但不把其描述为已完整上线功能。",
    )
    add_para_before(
        doc,
        marker,
        "界面设计方面，写作页强调编辑器稳定性和自动保存反馈；浏览页强调分类、搜索和卡片摘要；详情页强调正文阅读、锚点评论和互动按钮位置；个人中心后续承接收藏、合集和经验等级展示。各页面均应复用统一请求封装和错误提示组件，避免同类异常在不同页面表现不一致。",
    )
    add_table_before(
        doc,
        marker,
        "表6-4　界面验收要点",
        ["页面", "主要控件", "验收标准"],
        [
            ("写作页", "标题输入、Markdown 编辑器、AI 润色按钮、发布按钮。", "编辑过程不丢稿；发布失败有明确提示；AI 不可用时仍能保存和发布。"),
            ("浏览页", "分类筛选、关键词搜索、排序、笔记卡片。", "筛选条件变化后结果稳定；卡片信息足以判断是否进入详情。"),
            ("详情页", "正文渲染、摘要、评论区、互动按钮、合集导航。", "Markdown 展示正常；评论和表态后状态刷新。"),
            ("个人中心", "我的收藏、我的合集、经验记录。", "后续入口与后端数据结构一致，不重复定义用户资料。"),
        ],
    )
    insert_source_figure(doc, src, marker, "图 6-8 草稿发布与 AI 摘要降级顺序图")
    insert_source_figure(doc, src, marker, "图 6-11 笔记生命周期状态图")
    add_para_before(
        doc,
        marker,
        "关键算法并不追求复杂，而强调规则清晰。赞踩互斥通过“写入当前操作前移除对立操作”实现；收藏通过唯一约束保证重复点击不会产生重复记录；合集通过 note_id 唯一约束保证单归属；签到通过用户与日期的复合唯一键保证每日一次。这些规则均适合在服务层集中维护，并通过测试直接验证。",
    )
    add_para_before(
        doc,
        marker,
        "从实现风险看，最需要控制的是状态不同步问题。前端在用户点击点赞、收藏、发表评论后可以做乐观更新，但最终必须以服务端返回为准；若接口返回冲突或权限错误，前端应回滚本地状态并提示用户。对草稿自动保存，前端应展示保存中、已保存、保存失败三种状态，避免用户误以为内容已经入库。对合集导航，详情页应在上下文为空时隐藏前后篇入口，而不是展示不可点击的空链接。",
    )
    add_para_before(
        doc,
        marker,
        "详细设计中还需特别处理删除语义。笔记删除应优先采用软删除或状态变更，避免直接删除后造成评论、收藏、合集条目和外部引用失效；评论删除可以根据课程项目复杂度采用硬删或状态隐藏，但必须保证楼中楼回复不会变成孤儿数据。合集删除时只删除合集和条目关系，不删除笔记本身，因为笔记仍然属于作者的公开内容。",
    )
    add_para_before(
        doc,
        marker,
        "AI 相关接口的详细设计以“辅助、可拒绝、可降级”为原则。润色结果应以差异对照或候选文本方式返回，由用户决定是否采纳；摘要生成失败时可以使用本地截断摘要；未来审查分类只能作为管理员参考，不能在没有人工确认的情况下直接删除用户内容。该原则既符合课程文档中的可靠性要求，也避免模型误判造成内容治理风险。",
    )


def write_ch7(doc, marker):
    add_para_before(doc, marker, "1．笔记系统与社区互动模块（陶语涵）", "Heading 2")
    add_para_before(
        doc,
        marker,
        "测试范围覆盖当前已经实现或可直接验证的功能，包括草稿发布、信息流浏览、评论、赞踩互斥、收藏、合集和签到经验。规划功能只记录风险，不纳入通过项。",
    )
    add_para_before(
        doc,
        marker,
        "测试方法采用自动化测试与人工交互验证结合。后端使用 pytest 和内存 SQLite 验证服务函数、接口状态码和数据库约束；前端使用组件或纯函数测试验证表单状态、交互按钮和解析逻辑；人工测试用于检查富文本渲染、Markdown 展示、长评论分页和浏览器中的视觉反馈。",
    )
    add_para_before(
        doc,
        marker,
        "测试数据分为三类：一是普通公开笔记，用于浏览、详情和评论测试；二是当前用户私有草稿，用于自动保存和发布测试；三是边界数据，包括空标题、超长正文、重复收藏、重复签到、跨用户删除评论和合集冲突等。通过这三类数据可以覆盖正常路径、权限失败和数据一致性失败三种场景。",
    )
    add_table_before(
        doc,
        marker,
        "表7-1　笔记系统与社区互动模块测试用例",
        ["编号", "测试目标", "操作", "预期结果"],
        [
            ("TC-NOTE-01", "发布笔记", "作者从草稿页提交标题、正文、分类并发布。", "生成 visible 笔记；摘要失败时仍能发布。"),
            ("TC-NOTE-02", "评论层级", "对笔记发表评论，再对顶层评论回复，并尝试第三级回复。", "两层回复成功；第三级被拒绝。"),
            ("TC-NOTE-03", "赞踩互斥", "同一用户先点赞再点踩。", "点赞取消，点踩生效，计数保持一致。"),
            ("TC-NOTE-04", "收藏幂等", "重复收藏和取消同一篇笔记。", "不会产生重复记录，最终状态与最后一次操作一致。"),
            ("TC-NOTE-05", "合集归属", "作者把本人笔记加入合集，再尝试加入另一合集。", "首次成功；第二次因单归属约束失败。"),
            ("TC-NOTE-06", "每日签到", "同一用户同一天连续签到两次。", "第一次增加经验，第二次幂等返回且不重复发放。"),
        ],
    )
    add_para_before(
        doc,
        marker,
        "测试结论：模块核心链路能够支撑内容发布、公开浏览和基础社区互动。主要限制在于部分治理能力和前端入口仍需补齐，例如举报裁决、AI 审查、拉黑、等级徽章展示和合集拖拽排序；这些内容应作为后续迭代任务，而不影响本次课程设计对核心功能的验收。",
    )
    add_table_before(
        doc,
        marker,
        "表7-2　缺陷与后续建议",
        ["类别", "当前状态", "建议"],
        [
            ("前端入口", "部分收藏、合集、等级展示入口仍需完善。", "按用户高频路径优先补齐详情页和个人中心入口。"),
            ("治理能力", "举报、AI 审查和人工裁决为设计项。", "先实现举报记录和管理员处理，再接入 AI 辅助分类。"),
            ("性能风险", "评论和信息流随数据增长需要分页稳定性。", "保持游标分页，补充大数据量回归测试。"),
            ("一致性风险", "互动计数依赖后端记录汇总。", "所有写操作后重新读取服务端结果，避免前端本地计数漂移。"),
        ],
    )
    add_para_before(
        doc,
        marker,
        "从测试覆盖角度看，本模块应避免只验证“页面能打开”。课程验收时应重点演示五个闭环：草稿保存并发布、匿名浏览公开笔记、登录用户发表评论、同一用户赞踩互斥、作者把笔记加入合集并在详情页看到前后篇。上述闭环能证明模块已经具备从内容生产到内容消费的基本能力。",
    )
    add_para_before(
        doc,
        marker,
        "测试资源消耗主要为本地 SQLite 测试库、pytest/Vitest 单元测试、浏览器人工操作记录和少量示例笔记数据。测试过程中不依赖真实个人隐私材料，也不要求外部 AI 服务稳定在线。",
    )
    add_para_before(
        doc,
        marker,
        "最终评价为：本模块的核心功能已经覆盖内容平台最基本的生产、展示和互动需求，能够与资料库、导师库、会议库和班级空间共同构成飞跃手册的统一入口。后续若继续完善，应优先补齐前端入口和治理闭环，再考虑更复杂的推荐、热榜和等级权益，以免在课程项目阶段引入过高实现风险。",
    )
    add_para_before(
        doc,
        marker,
        "测试结论还表明，笔记模块的质量关键不在单个接口是否能返回 200，而在多接口组合后的状态是否一致。例如发布后应能在浏览页检索到，在详情页能发表评论，收藏后能在个人入口读取到，加入合集后能看到前后篇上下文。后续回归测试应围绕这些组合路径组织，而不是只保留孤立的接口单测。",
    )


def remove_empty_paragraphs(doc):
    body = doc._body._element
    removed = 0
    for p in list(doc.paragraphs):
        if p.text.strip():
            continue
        xml = p._element.xml
        if "w:drawing" in xml or "w:pict" in xml or "w:br" in xml:
            continue
        body.remove(p._element)
        removed += 1
    return removed


def main():
    backup = TARGET.with_name(f"{TARGET.stem}.balanced-bak-{datetime.now():%Y%m%d-%H%M%S}{TARGET.suffix}")
    shutil.copy2(TARGET, backup)
    doc = Document(str(TARGET))
    replace_range(doc, "四", write_ch4)
    replace_range(doc, "五", write_ch5)
    replace_range(doc, "六", write_ch6)
    replace_range(doc, "七", write_ch7)
    removed = remove_empty_paragraphs(doc)
    doc.save(str(TARGET))
    print(f"balanced: {TARGET}")
    print(f"backup: {backup}")
    print(f"removed empty paragraphs: {removed}")


if __name__ == "__main__":
    main()
