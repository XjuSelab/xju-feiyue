from __future__ import annotations

import copy
import shutil
from datetime import datetime
from io import BytesIO
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn


ROOT = Path(__file__).resolve().parents[1]
TARGET = ROOT / "docs/course-design/赵文彪-飞跃.docx"
SOURCE = ROOT / "赵文彪-飞跃-陶语涵.docx"

CHAPTERS = [
    ("四、系统需求分析", ("四", "4"), ("五", "5"), "1．笔记系统与社区互动模块（陶语涵）"),
    ("五、系统概要设计", ("五", "5"), ("六", "6"), "1．笔记系统与社区互动模块（陶语涵）"),
    ("六、系统详细设计", ("六", "6"), ("七", "7"), "1．笔记系统与社区互动模块（陶语涵）"),
    ("七、系统测试", ("七", "7"), ("八", "8"), "1．笔记系统与社区互动模块（陶语涵）"),
]


def heading_text(p) -> str:
    return "".join(p.text.split())


def is_heading(p, starts: tuple[str, ...], level: int | None = None) -> bool:
    text = heading_text(p)
    if not p.style or not p.style.name.startswith("Heading"):
        return False
    if level is not None and p.style.name != f"Heading {level}":
        return False
    return any(text.startswith(s) for s in starts)


def para_by_element(doc: Document):
    return {p._element: p for p in doc.paragraphs}


def source_section_elements(doc: Document, start_keys: tuple[str, ...], end_keys: tuple[str, ...]):
    body = doc._body._element
    children = list(body)
    pmap = para_by_element(doc)
    start = end = None
    for i, child in enumerate(children):
        p = pmap.get(child)
        if p is None:
            continue
        if start is None and is_heading(p, start_keys, level=1):
            start = i
            continue
        if start is not None and is_heading(p, end_keys, level=1):
            end = i
            break
    if start is None:
        raise RuntimeError(f"未找到陶语涵文档章节起点 {start_keys}")
    if end is None:
        end = len(children)
    return children[start + 1 : end]


def remap_image_relationships(src_doc: Document, dst_doc: Document, element) -> None:
    for node in element.xpath(".//*[@r:embed]"):
        old = node.get(qn("r:embed"))
        part = src_doc.part.related_parts.get(old)
        if part is None or not hasattr(part, "blob"):
            continue
        new_rid, _ = dst_doc.part.get_or_add_image(BytesIO(part.blob))
        node.set(qn("r:embed"), new_rid)


def demote_heading_styles(element) -> None:
    for p in element.xpath(".//w:p"):
        pstyle = p.find(qn("w:pPr") + "/" + qn("w:pStyle"))
        if pstyle is None:
            continue
        val = pstyle.get(qn("w:val"))
        if val == "Heading2":
            pstyle.set(qn("w:val"), "Heading3")
        elif val == "Heading3":
            pstyle.set(qn("w:val"), "Heading4")
        elif val == "Heading4":
            pstyle.set(qn("w:val"), "Normal")


def add_para_before(doc: Document, marker, text: str, style: str, align=None):
    p = doc.add_paragraph(text, style=style)
    if align is not None:
        p.alignment = align
    el = p._element
    doc._body._element.remove(el)
    marker.addprevious(el)
    return el


def find_placeholder_range(doc: Document, chapter_keys: tuple[str, ...]):
    children = list(doc._body._element)
    pmap = para_by_element(doc)
    chapter_seen = False
    start = end = None
    for i, child in enumerate(children):
        p = pmap.get(child)
        if p is None:
            continue
        if not chapter_seen and is_heading(p, chapter_keys, level=1):
            chapter_seen = True
            continue
        if not chapter_seen:
            continue
        if p.style.name == "Heading 1":
            break
        text = heading_text(p)
        if start is None and p.style.name == "Heading 2" and "陶语涵" in text:
            start = i
            continue
        if start is not None and (p.style.name in {"Heading 1", "Heading 2"}):
            end = i
            break
    if start is None:
        raise RuntimeError(f"未找到总文档陶语涵占位 {chapter_keys}")
    if end is None:
        end = len(children)
    return start, end


def replace_placeholder_with_section(dst: Document, src: Document, start_keys, end_keys, module_heading: str) -> None:
    body = dst._body._element
    children = list(body)
    start, end = find_placeholder_range(dst, start_keys)
    marker = children[end]
    for child in children[start:end]:
        body.remove(child)

    add_para_before(dst, marker, module_heading, style="Heading 2")
    for src_el in source_section_elements(src, start_keys, end_keys):
        el = copy.deepcopy(src_el)
        remap_image_relationships(src, dst, el)
        demote_heading_styles(el)
        marker.addprevious(el)


def refresh_overview_text(doc: Document) -> None:
    replacements = {
        "当前已完成资料库与教务学分自查、导师/院校库与 CCF 会议截稿、用户班级小组权限与任务管理三个模块；笔记系统由陶语涵负责，因尚未提交材料，本合并稿保留空位。":
            "当前已汇总笔记系统与社区互动、资料库与教务学分自查、导师/院校库与 CCF 会议截稿、用户班级小组权限与任务管理四个模块。",
        "陶语涵模块对应笔记写作、浏览、评论和点赞，当前保留待补。":
            "陶语涵模块对应笔记写作、浏览、评论、点赞、合集、签到与社区互动。",
    }
    for para in doc.paragraphs:
        text = para.text
        new_text = text
        for old, new in replacements.items():
            new_text = new_text.replace(old, new)
        if new_text != text:
            para.clear()
            para.add_run(new_text)


def main() -> None:
    backup = TARGET.with_name(f"{TARGET.stem}.tao-bak-{datetime.now():%Y%m%d-%H%M%S}{TARGET.suffix}")
    shutil.copy2(TARGET, backup)
    dst = Document(str(TARGET))
    src = Document(str(SOURCE))

    refresh_overview_text(dst)
    for _title, start_keys, end_keys, module_heading in CHAPTERS:
        replace_placeholder_with_section(dst, src, start_keys, end_keys, module_heading)

    dst.save(str(TARGET))
    print(f"merged tao: {TARGET}")
    print(f"backup: {backup}")


if __name__ == "__main__":
    main()
